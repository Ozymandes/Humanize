#!/usr/bin/env python3
"""
Text Humanizer Web App
A beautiful web interface for humanizing AI-generated text
Supports PDF, Word (.docx), and plain text files
"""

from flask import Flask, render_template, request, send_file, after_this_request
from werkzeug.utils import secure_filename
import os
from pathlib import Path
import tempfile
from noisy_humanize import noisy_humanize, clean_humanized_text
import docx
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import fitz  # PyMuPDF
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
app = Flask(__name__, template_folder=template_dir)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.config['TEMPLATES_AUTO_RELOAD'] = True

# Force clear Jinja cache
app.jinja_env.cache = {}
app.jinja_env.auto_reload = True

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    # Direct file read to bypass any template caching
    template_path = os.path.join(template_dir, 'index.html')
    with open(template_path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_text_from_file(file, file_ext):
    """Extract text from different file formats while preserving structure"""
    if file_ext == 'txt':
        return file.read().decode('utf-8')

    elif file_ext == 'docx':
        # Save the original file for later reference
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_docx.docx')
        with open(temp_path, 'wb') as f:
            f.write(file.read())
        return temp_path  # Return path to preserve formatting

    elif file_ext == 'pdf':
        # Save the original file for later reference
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_pdf.pdf')
        with open(temp_path, 'wb') as f:
            f.write(file.read())

        # Also extract text for humanization
        doc = fitz.open(temp_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return {'text': '\n\n'.join(text_parts), 'path': temp_path}

    return ''

def create_output_file(text, file_ext, original_filename, original_file_info=None):
    """Create output file in the same format as input, preserving formatting"""
    output_filename = Path(original_filename).stem + '_humanized'

    if file_ext == 'txt':
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(output_filename + '.txt'))
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return output_path, 'text/plain', output_filename + '.txt'

    elif file_ext == 'docx' and original_file_info:
        # Preserve formatting from original docx
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(output_filename + '.docx'))
        original_doc = docx.Document(original_file_info)
        new_doc = docx.Document()

        # Copy styles and formatting from original
        # Split humanized text into paragraphs and apply original formatting
        humanized_paras = text.split('\n\n')

        for i, para in enumerate(original_doc.paragraphs):
            if i < len(humanized_paras) and humanized_paras[i].strip():
                # Create new paragraph with humanized text but original formatting
                new_para = new_doc.add_paragraph(humanized_paras[i].strip())

                # Copy paragraph formatting
                new_para.style = para.style
                new_para.alignment = para.alignment
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing

        new_doc.save(output_path)

        # Clean up temp file
        try:
            os.remove(original_file_info)
        except:
            pass

        return output_path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', output_filename + '.docx'

    elif file_ext == 'pdf' and original_file_info:
        # Preserve formatting from original PDF
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(output_filename + '.pdf'))

        # Open original PDF to get formatting
        original_doc = fitz.open(original_file_info)

        # Create new PDF with same page dimensions
        new_doc = fitz.open()  # Create new PDF

        # Get humanized paragraphs
        humanized_paras = text.split('\n\n')

        para_index = 0
        for page_num in range(len(original_doc)):
            page = original_doc[page_num]
            rect = page.rect

            # Create new page with same dimensions
            new_page = new_doc.new_page(width=rect.width, height=rect.height)

            # Add humanized text preserving basic layout
            if para_index < len(humanized_paras):
                # Insert text with basic formatting
                text_rect = fitz.Rect(rect.x0 + 50, rect.y0 + 50, rect.x1 - 50, rect.y1 - 50)
                new_page.insert_textbox(
                    text_rect,
                    humanized_paras[para_index].strip(),
                    fontsize=11,
                    fontname="helv",
                    align=0
                )
                para_index += 1

        new_doc.save(output_path)
        original_doc.close()
        new_doc.close()

        # Clean up temp file
        try:
            os.remove(original_file_info)
        except:
            pass

        return output_path, 'application/pdf', output_filename + '.pdf'

    return None, None, None

@app.route('/humanize', methods=['POST'])
def humanize():
    if 'file' not in request.files:
        return {'error': 'No file provided'}, 400

    file = request.files['file']
    if file.filename == '':
        return {'error': 'No file selected'}, 400

    if not allowed_file(file.filename):
        return {'error': 'Only .txt, .pdf, and .docx files are allowed'}, 400

    try:
        # Get file extension
        file_ext = Path(file.filename).suffix.lstrip('.').lower()

        # Extract text from file (preserving original file info)
        file_data = extract_text_from_file(file, file_ext)

        # Handle different return types from extract_text_from_file
        if file_ext == 'txt':
            text = file_data
            original_file_info = None
        elif file_ext == 'docx':
            original_file_info = file_data
            # Extract text from the saved file
            doc = docx.Document(original_file_info)
            paragraphs = [para.text for para in doc.paragraphs]
            text = '\n\n'.join(paragraphs)
        elif file_ext == 'pdf':
            text = file_data['text']
            original_file_info = file_data['path']
        else:
            text = ''
            original_file_info = None

        if not text.strip():
            return {'error': 'No text found in file'}, 400

        # Get optional seed parameter
        seed = request.form.get('seed')
        seed_value = int(seed) if seed and seed.isdigit() else None

        # Humanize the text
        humanized_text = noisy_humanize(text, seed=seed_value)

        # Create output file in the same format, preserving formatting
        output_path, mimetype, output_filename = create_output_file(
            humanized_text,
            file_ext,
            file.filename,
            original_file_info
        )

        # Clean up the file after sending
        @after_this_request
        def remove_file(response):
            try:
                if output_path and os.path.exists(output_path):
                    os.remove(output_path)
            except Exception as e:
                app.logger.error(f'Error removing file: {e}')
            return response

        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype=mimetype
        )

    except Exception as e:
        app.logger.error(f'Error processing file: {e}')
        return {'error': f'Error processing file: {str(e)}'}, 500

@app.route('/humanize-text', methods=['POST'])
def humanize_text():
    """Endpoint for direct text input (no file upload)"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return {'error': 'No text provided'}, 400

        text = data['text']
        if not text.strip():
            return {'error': 'Text cannot be empty'}, 400

        # Get optional parameters
        seed = data.get('seed')
        seed_value = int(seed) if seed and seed.isdigit() else None
        intensity = data.get('intensity', 'intermediate')

        # Humanize the text
        humanized_text = noisy_humanize(text, seed=seed_value, intensity=intensity)

        return {'humanized_text': humanized_text}

    except Exception as e:
        app.logger.error(f'Error processing text: {e}')
        return {'error': f'Error processing text: {str(e)}'}, 500

@app.route('/clean-text', methods=['POST'])
def clean_text():
    """Remove filler words and personal asides from humanized text"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return {'error': 'No text provided'}, 400

        text = data['text']
        if not text.strip():
            return {'error': 'Text cannot be empty'}, 400

        # Clean the text
        cleaned_text = clean_humanized_text(text)

        return {'cleaned_text': cleaned_text}

    except Exception as e:
        app.logger.error(f'Error cleaning text: {e}')
        return {'error': f'Error cleaning text: {str(e)}'}, 500

if __name__ == '__main__':
    # Create templates directory if it doesn't exist
    templates_dir = Path(__file__).parent / 'templates'
    templates_dir.mkdir(exist_ok=True)

    # Get port from environment variable (for cloud deployment)
    port = int(os.environ.get('PORT', 8000))  # Changed to 8000 to bypass caching

    # Run the app
    app.run(debug=True, host='0.0.0.0', port=port, use_reloader=False)
